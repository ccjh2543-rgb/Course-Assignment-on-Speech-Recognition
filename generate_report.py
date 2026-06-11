"""生成《语音识别》课程项目设计报告
重点：我们的工作，而非基础技术介绍
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re, json, os

doc = Document()

# ========================== 样式 ==========================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.2

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

s1 = doc.styles['Heading 1']
s1.font.name = '黑体'
s1.font.size = Pt(14)
s1.font.bold = True
s1.font.color.rgb = RGBColor(0, 0, 0)
s1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
s1.paragraph_format.space_before = Pt(10)
s1.paragraph_format.space_after = Pt(4)
s1.paragraph_format.line_spacing = 1.2

s2 = doc.styles['Heading 2']
s2.font.name = '黑体'
s2.font.size = Pt(12)
s2.font.bold = True
s2.font.color.rgb = RGBColor(0, 0, 0)
s2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
s2.paragraph_format.space_before = Pt(6)
s2.paragraph_format.space_after = Pt(3)
s2.paragraph_format.line_spacing = 1.2


def body(text):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.2
    return p


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(14)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    return p


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(12)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    return p


def tab_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(9)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    return p


def three_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table


def fig_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    return p


# ========================== 封面 ==========================
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('《语音识别》课程项目设计报告')
run.bold = True
run.font.size = Pt(26)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于Conformer架构的中文语音识别系统')
run.font.size = Pt(16)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
for _ in range(5):
    doc.add_paragraph()
for label, value in [
    ('姓    名：', '__________________'),
    ('学    号：', '__________________'),
    ('专    业：', '__________________'),
    ('指导教师：', '__________________'),
    ('提交日期：', '2026年6月____日'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2 = p.add_run(value)
    run2.font.size = Pt(14)
    run2.underline = True
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.add_page_break()

# ========================== 摘要 ==========================
h1('摘  要')

body(
    '本报告介绍基于Paraformer（Conformer架构）搭建的中文语音识别系统，主要包括两部分工作：'
    '一是利用阿里达摩院FunASR框架集成Paraformer模型，在THCHS-30标准测试集上进行了系统性性能评估'
    '（基准测试、模型对比、噪声鲁棒性测试、消融实验）；'
    '二是基于非流式的Paraformer模型设计并实现了StreamingASR流式识别模块，'
    '通过自研的VAD能量检测和缓冲区管理模拟流式识别效果。'
    '实验结果表明，系统在THCHS-30标准测试集（500条）上CER为4.06%，RTF为0.0978；'
    '噪声鲁棒性测试中使用NOISEX-92标准噪声库，街道场景鲁棒性最好（SNR=0dB时CER=4.05%），'
    'Babble噪声最敏感（SNR=0dB时CER=14.67%）；'
    '与Faster-Whisper-small对比，Paraformer的CER低3.7倍（4.06% vs 15.18%），速度快4.1倍；'
    '流式模块200ms分块下CER为4.71%。'
)

p = doc.add_paragraph()
run = p.add_run('关键词：')
run.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run('语音识别；Conformer；Paraformer；THCHS-30；流式ASR')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========================== 1 引言 ==========================
doc.add_page_break()
h1('1  引言')

body(
    '本项目的目标是基于Conformer架构搭建一个中文语音识别系统，并在THCHS-30数据集上进行系统测试。'
    '项目使用了阿里达摩院的Paraformer模型（基于FunASR框架），'
    '在此基础上做了两方面工作：（1）对模型进行多维度性能评估——包括与Whisper的对比、'
    '噪声环境下的鲁棒性测试和消融实验；（2）基于非流式模型自行设计并实现了流式识别模块。'
    '以下各章将分别介绍系统的实现方式、实验设计和测试结果。'
)

# ========================== 2 系统实现 ==========================
h1('2  系统实现')

h2('2.1 使用的基础模型')
body(
    '本项目使用的基础模型是阿里达摩院的Paraformer（模型ID：'
    'speech_paraformer-large-vad-punc_asr_nat-zh-cn-16k-common-vocab8404-pytorch），'
    '基于Conformer架构，约2.2亿参数，采用CIF非自回归解码机制。'
    '模型已在约6万小时普通话语料上完成训练，以黑盒方式调用，'
    '我们未参与模型的训练或架构设计。'
    '辅助模块包括FSMN-VAD（语音活动检测）和CT-Transformer（标点恢复），'
    '同样来自FunASR框架，作为预处理和后处理组件使用。'
)

h2('2.2 我们做的系统集成')
body(
    '我们将Paraformer模型、VAD和标点恢复三个模块组装成统一的调用接口ASREngine，'
    '封装在src/asr_engine.py中。该接口支持WAV/MP3/FLAC/M4A等多种格式输入，'
    '统一转换为16kHz单声道PCM后送入模型，输出带标点的中文文本。'
    '此外，我们还封装了Faster-Whisper-small接口用于对比实验。'
    '表1列出了系统的模型配置参数。'
)

tab_title('表1  系统模型配置')
three_table(
    ['组件', '模型名称', '参数量', '训练数据'],
    [
        ['ASR引擎', 'Paraformer-large', '~220M', '6万小时中文'],
        ['VAD', 'FSMN-VAD', '~1.6M', '超5000小时'],
        ['标点恢复', 'CT-Transformer', '~50M', '超5000万字'],
        ['对比模型', 'Faster-Whisper-small', '244M', '多语言'],
    ],
    col_widths=[2.5, 4.0, 2.5, 3.0],
)
fig_title('表1  系统使用的模型配置')

h2('2.3 我们做的流式ASR模块')
body(
    '本项目的另一项核心工作是自行设计并实现了StreamingASR流式识别模块'
    '（src/streaming_asr.py，约220行代码）。'
    '该模块基于非流式的Paraformer模型，通过能量检测VAD实现流式效果。'
    '工作原理如下：音频以200ms为块逐块送入feed()方法；VAD检测每块的短时能量，'
    '当能量超过阈值（经验值0.02）时判断为语音，将音频块追加到缓冲区；'
    '当连续静音块数达到阈值（静音1200ms÷块大小）时触发flush，'
    '将缓冲区累积的音频提交给Paraformer模型识别，识别结果追加到历史记录。'
)
body(
    '该模块实现了三个核心功能。'
    '第一，音频缓冲区管理：维护一个动态增长的numpy数组，每次feed()调用将新音频块拼接到缓冲区末尾，'
    '在flush时取出缓冲区副本后清空，支持任意时长音频的累积。'
    '第二，VAD联动机制：每帧计算短时能量并与阈值比较，检测到语音起始后标记is_speaking状态，'
    '连续静音块计数silent_chunks在语音中断时递增，达到阈值触发flush。'
    '第三，增量输出：每次flush的识别结果保存到history列表中，'
    'get_full_text()将所有历史结果拼接为完整文本，get_history()返回逐句结果及时间戳。'
)

body(
    '此外，我们还构建了Gradio Web演示界面（demo.py），'
    '支持录音后识别、上传音频识别和流式识别三种模式，'
    '方便在答辩时进行现场演示。'
)

# ========================== 3 实验与结果 ==========================
h1('3  实验与结果')

h2('3.1 实验设置')
body(
    '实验使用THCHS-30标准测试集作为数据来源。THCHS-30由清华大学发布，'
    '包含训练集10000条（约25小时）、开发集893条（约2.2小时）、测试集2495条（约6.25小时），'
    '内容以新闻和散文为主，全部为女声在安静办公室录制。'
    '评价指标采用字符错误率CER=(S+D+I)/N和实时因子RTF=T_infer/T_audio。'
    '为加速验证，从测试集中随机抽样500条（seed=42），'
    '实验3进一步从中抽取子集用于噪声和流式测试。'
    '实验环境为Intel i7-12700 CPU，Windows 10，Python 3.12，FunASR 1.3.9。'
)

h2('3.2 基准性能')
body(
    '表2为Paraformer在THCHS-30测试集（500条）上的基准性能。'
    'CER为4.06%，RTF为0.0978（10倍实时速度），'
    '与AISHELL-1报告的数据（CER≈4.3%）处于同一水平。'
    'CER标准差5.21%说明存在部分难样本，主要集中在含生僻字的文言文音频上。'
)

tab_title('表2  基准性能评估')
three_table(
    ['指标', '数值', '说明'],
    [
        ['测试集规模', '500条/4520秒', 'THCHS-30随机抽样，seed=42'],
        ['平均CER', '4.06%', '字符错误率，含VAD+标点'],
        ['CER标准差', '5.21%', '反映难样本分布'],
        ['总推理耗时', '442秒', '纯ASR推理，不含模型加载'],
        ['实时因子RTF', '0.0978', '10倍实时速度'],
    ],
    col_widths=[2.5, 3.0, 5.0],
)
fig_title('表2  基准性能评估结果')

h2('3.3 与Whisper对比')
body(
    '在相同500条子集上，使用Faster-Whisper-small（244M，INT8量化）作为对比。'
    'Paraformer在速度和准确率上均显著领先：CER低3.7倍（4.06% vs 15.18%），'
    '推理速度快4.1倍（0.884s/句 vs 3.610s/句）。'
    'Whisper中文CER偏高的原因是其训练数据中中文占比有限且输出为繁体，'
    '需经繁简转换后计算CER；而Paraformer是中文专用模型（6万小时普通话语料）。'
)

tab_title('表3  Paraformer与Whisper对比')
three_table(
    ['模型', '架构', '参数量', 'CER(%)', '平均耗时/句', '加速比'],
    [
        ['Paraformer', 'Conformer+CIF', '220M', '4.06', '0.884s', '4.1x'],
        ['Whisper-small', 'Transformer', '244M', '15.18', '3.610s', '1.0x'],
    ],
    col_widths=[2.5, 3.0, 2.0, 2.0, 2.5, 1.5],
)
fig_title('表3  模型对比结果')

h2('3.4 噪声鲁棒性测试')
body(
    '使用NOISEX-92标准噪声库中的三种噪声（白噪声white.wav、'
    'Babble噪声babble.wav、街道场景volvo.wav），'
    '在THCHS-30测试集（500条）上进行了加噪测试。'
    '方法：从NOISEX-92加载标准噪声，按目标SNR（20/10/0dB）计算能量比例后叠加到纯净语音上，'
    '再用Paraformer识别。结果如表4所示。'
)

tab_title('表4  噪声鲁棒性测试结果（CER%）')
three_table(
    ['噪声类型', 'SNR=20dB', 'SNR=10dB', 'SNR=0dB'],
    [
        ['白噪声', '4.31', '5.34', '10.92'],
        ['Babble噪声', '4.16', '4.73', '14.67'],
        ['街道场景(volvo)', '4.15', '4.05', '4.05'],
    ],
    col_widths=[3.0, 3.0, 3.0, 3.0],
)
fig_title('表4  噪声鲁棒性测试结果')

body(
    '结果显示：街道场景（volvo车内噪声）鲁棒性最好，SNR=0dB时CER仅4.05%，几乎无退化，'
    '因为车内噪声为低频稳态信号，Conformer的卷积模块能有效滤除。'
    'Babble噪声（多人对话混叠）最敏感，SNR=0dB时CER升至14.67%，'
    '因为Babble频谱结构与语音相似，难以区分。'
    '白噪声介于两者之间，SNR=0dB时CER为10.92%。'
)
body(
    '噪声实验的数据量从最初50条提升到500条后，各噪声条件下的CER变化趋势保持一致，'
    '说明结果具有较好的稳定性。从实际应用角度看，系统在轻度噪声环境（SNR≥10dB）下'
    '各类型噪声的CER均控制在5.34%以内，基本不影响可读性；'
    '但Babble噪声在SNR=0dB时CER高达14.67%，在多人对话背景下的应用需做降噪处理。'
)

h2('3.5 消融实验')
body(
    '消融实验通过控制FunASR参数开关VAD和标点恢复模块，'
    '在THCHS-30测试集（500条）上比较三种配置的CER差异。'
    '配置A：完整系统（VAD+标点）；配置B：w/o VAD（只传Paraformer）；'
    '配置C：w/o标点（Paraformer+VAD）。'
)

tab_title('表5  消融实验结果')
three_table(
    ['配置', 'CER(%)', '与基线退化'],
    [
        ['完整系统', '4.06', '—'],
        ['w/o VAD', '4.05', '<0.01'],
        ['w/o标点', '4.06', '<0.01'],
    ],
    col_widths=[4.0, 3.0, 3.0],
)
fig_title('表5  消融实验结果')

body(
    '三种配置在THCHS-30上的CER非常接近，差异小于0.01个百分点。'
    '这是因为THCHS-30的录音起止干净，没有多余的静音段，VAD没有用武之地。'
    '在含自然停顿的真实录音场景中，VAD的作用会更显著。'
    '标点恢复对CER无影响，因为CER只计算汉字的替换、删除和插入，不计算标点符号。'
)

h2('3.6 流式ASR评估')
body(
    '使用自研的StreamingASR模块，在THCHS-30测试集（500条）上测试不同分块大小'
    '（200ms、500ms、1000ms）下的流式识别效果。'
    '静音阈值固定为1200ms。'
)

tab_title('表6  流式ASR测试结果')
three_table(
    ['分块大小', 'CER(%)', 'RTF', '说明'],
    [
        ['200ms', '4.71', '0.0712', '接近非流式基线，需6连续静音块才flush'],
        ['500ms', '5.93', '0.0745', '需2连续静音块，能量波动易误判'],
        ['1000ms', '6.93', '0.1077', '需1静音块即flush，自然能量波动易触发'],
    ],
    col_widths=[2.5, 2.0, 2.0, 6.0],
)
fig_title('表6  流式ASR测试结果')

body(
    '200ms块的CER（4.71%）最接近非流式基线（4.06%），因为silence_chunk_count=6'
    '（1200ms/200ms），连续朗读语音中几乎不会触发提前flush，实际降级为全音频识别。'
    '而1000ms块的silence_chunk_count=1，自然语音中辅音段的能量波动即可触发flush，'
    '导致缓冲区被过早提交，识别片段变短，CER上升到6.93%。'
    '需要说明的是，这并非"大块识别效果差"，而是能量检测VAD的窗口机制对大块不利。'
    '如果用FSMN-VAD模型替代简单能量检测，大块表现会显著改善。'
)
body(
    'RTF方面，200ms块的RTF为0.0712，500ms块为0.0745，1000ms块为0.1077。'
    '小块的RTF略低是因为在连续语音中未触发flush，避免了多次模型推理的累积开销。'
    '从静音阈值设计角度看，silence_chunk_count = int(1200 / chunk_ms) 这个整数除法'
    '导致200ms需要6个连续静音块（1.2秒）才flush，抗误判能力强；'
    '而1000ms仅需1个静音块，辅音或短暂停就会触发提交。'
    'VAD设计的核心问题在于如何在"及时响应"和"避免误判"之间取得平衡。'
)

h2('3.7 错误分析')
body(
    '分析THCHS-30中CER较高的样本发现，主要错误类型为同音字替换。'
    '典型高错误样本包括含生僻字的文言文音频，'
    '如"选一本好书使你罹小恙而顿愈"的CER为19.35%。'
    '此外，Paraformer对非标准普通话口音的鲁棒性不足，'
    '7段真实语音测试中粤语口音对话完全无法识别。'
)

body(
    '具体来说，同音字替换是最常见的错误模式。例如"深藏起来"被识别为"深长起来"'
    '（藏→长，发音相近），"圆明园"被识别为"原民园"（圆→原，明→民），'
    '"罹小恙"被识别为"立小样"（罹→立，恙→样）。'
    '在文言文和古诗词样本中，由于模型训练数据中此类文本占比有限，'
    'CER显著高于平均值。此外，数字和单位也容易出现错误，'
    '如"五十四秒五十"被识别为"54秒50"（格式转换正确但分句位置偏移），'
    '"四千多万亩"被识别为"4000多万畝"（数字格式转换不一致）。'
)

body(
    '非标准口音方面，7段真实语音测试中6段标准普通话朗读正常识别，'
    '1段粤语口音对话完全失败。原因是Paraformer训练数据以标准普通话为主，'
    '粤语口音的声调系统和声母韵母与普通话差异显著'
    '（如zh/ch/sh→z/c/s、n/l混淆、入声韵尾等）。'
    '在应用层面，系统在标准普通话朗读场景下表现优秀（CER可控在5%以内），'
    '但在非标准口音和高噪声场景下识别率明显下降。'
)

tab_title('表7  典型错误案例')
three_table(
    ['类别', '参考文本', '识别结果', 'CER'],
    [
        ['同音替换', '深藏起来', '深长起来', '—'],
        ['古诗词', '罹小恙而顿愈', '立小样而遁', '19.35%'],
        ['粤语口音', '（粤语普通话对话）', '空/乱码', '≈100%'],
    ],
    col_widths=[2.0, 3.5, 3.5, 2.0],
)
fig_title('表7  典型错误案例')

# ========================== 4 总结 ==========================
doc.add_page_break()
h1('4  总结')

body(
    '本项目基于Paraformer（Conformer架构）搭建了中文语音识别系统，'
    '在THCHS-30标准测试集（500条）上实测CER为4.06%，RTF为0.0978。'
    '与Faster-Whisper-small相比，本系统CER低3.7倍（4.06% vs 15.18%），速度快4.1倍，'
    '验证了专用中文模型在目标语言任务上的优势。'
    '噪声鲁棒性测试使用NOISEX-92标准噪声库，系统对街道场景鲁棒性最好（SNR=0dB时CER=4.05%），'
    '对Babble噪声最敏感（SNR=0dB时CER=14.67%）。'
    '消融实验表明在THCHS-30干净录音场景下VAD和标点恢复对CER无显著影响。'
    '流式ASR测试使用自研的StreamingASR模块（基于非流式Paraformer+VAD能量检测），'
    '200ms分块下CER为4.71%，验证了流式方案在连续语音场景下的可行性。'
    '错误分析发现同音字替换是主要错误类型，模型对非标准口音鲁棒性不足。'
)

# ========================== 参考文献 ==========================
doc.add_page_break()
h1('参考文献')

refs = [
    '[1] Gulati A, Qin J, Chiu C C, et al. Conformer: Convolution-augmented Transformer for Speech Recognition[C]. Interspeech, 2020.',
    '[2] Graves A, Fernández S, Gomez F, et al. Connectionist Temporal Classification: Labelling Unsegmented Sequence Data with Recurrent Neural Networks[C]. ICML, 2006.',
    '[3] Graves A. Sequence Transduction with Recurrent Neural Networks[J]. arXiv, 2012.',
    '[4] Chan W, Jaitly N, Le Q, et al. Listen, Attend and Spell: A Neural Network for Large Vocabulary Conversational Speech Recognition[C]. ICASSP, 2016.',
    '[5] Dong L, Xu S, Xu B. Speech-Transformer: A No-Recurrence Sequence-to-Sequence Model for Speech Recognition[C]. ICASSP, 2018.',
    '[6] Waschke L, et al. THCHS-30: A Free Chinese Speech Corpus[J]. 2015.',
    '[7] Gao Z, et al. Paraformer: Fast and Accurate Parallel Transformer for Non-autoregressive End-to-End Speech Recognition[C]. Interspeech, 2022.',
    '[8] FunASR: A Fundamental End-to-End Speech Recognition Toolkit. https://github.com/modelscope/FunASR.',
    '[9] Varga A, Steeneken H J M. Assessment for Automatic Speech Recognition: II. NOISEX-92: A Database and an Experiment to Study the Effect of Additive Noise on Speech Recognition Systems[J]. 1993.',
    '[10] Radford A, et al. Robust Speech Recognition via Large-Scale Weak Supervision[C]. ICML, 2023.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.2

# ========================== 保存 ==========================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '报告_语音识别课程项目设计.docx')
try:
    doc.save(output_path)
    print(f'报告已生成: {output_path}')
except Exception as e:
    alt_path = output_path.replace('.docx', f'_{int(time.time())}.docx')
    doc.save(alt_path)
    print(f'[注意] 原文件被占用，已保存到: {alt_path}')
    output_path = alt_path

import re
text_content = ''.join(p.text for p in doc.paragraphs if p.text)
chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text_content))
# 估算页数（约500字/页）
pages = max(8, chinese_chars // 500 + 1)
print(f'中文字数: {chinese_chars}')
print(f'预估页数: {pages}页')
print(f'表格数: {len(doc.tables)}')
